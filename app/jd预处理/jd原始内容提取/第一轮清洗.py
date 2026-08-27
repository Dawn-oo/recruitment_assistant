from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from pydantic import BaseModel, Field


# 1. Schema
class RawResponsibility(BaseModel):

    sequence: int

    description: str

    time_percentage: float | None = None

    tasks: list[str] = Field(
        default_factory=list
    )


class RawQualification(BaseModel):

    minimum_education: str | None = None

    education_background: str | None = None

    work_experience_raw: str | None = None

    training_subjects: list[str] = Field(
        default_factory=list
    )

    competencies: list[str] = Field(
        default_factory=list
    )


class RawJobDocument(BaseModel):

    source_file: str

    source_start_table: int
    source_start_row: int

    source_end_table: int
    source_end_row: int

    # 岗位基本信息

    job_title: str

    job_code: str | None = None

    department: str | None = None

    headcount: int | None = None

    direct_supervisor: str | None = None

    direct_reports: list[str] = Field(
        default_factory=list
    )

    job_family: str | None = None

    team_size: int | None = None

    salary_grade_raw: str | None = None

    analysis_date_raw: str | None = None

    # 岗位内容

    job_summary: str | None = None

    responsibilities: list[RawResponsibility] = Field(
        default_factory=list
    )

    permissions: list[str] = Field(
        default_factory=list
    )

    # 工作关系

    reporting_relationship: str | None = None

    internal_collaboration: str | None = None

    external_collaboration: str | None = None

    # 工作条件

    work_environment: str | None = None

    tools_and_equipment: list[str] = Field(
        default_factory=list
    )

    required_documents: list[str] = Field(
        default_factory=list
    )

    # 任职资格

    qualification: RawQualification = Field(
        default_factory=RawQualification
    )

    # 其他

    remarks: str | None = None

    raw_fields: dict[str, str] = Field(
        default_factory=dict
    )

    # 数据质量检查结果
    warnings: list[str] = Field(
        default_factory=list
    )

# 2. 字段别名

FIELD_ALIASES = {

    "job_title": [
        "岗位名称",
        "职位名称",
    ],

    "job_code": [
        "岗位编号",
        "职位编号",
    ],

    "department": [
        "所在部门",
        "所属部门",
    ],

    "headcount": [
        "岗位定员",
    ],

    "direct_supervisor": [
        "直接上级",
    ],

    "direct_reports": [
        "直接下属部门/岗位",
        "直接下属",
    ],

    "job_family": [
        "职系",
    ],

    "team_size": [
        "所辖人员数目",
        "所辖人数",
    ],

    "salary_grade_raw": [
        "工资级别范围",
        "薪资级别范围",
    ],

    "analysis_date_raw": [
        "岗位分析日期",
    ],

    "job_summary": [
        "本职概述",
        "岗位概述",
    ],

    "permissions": [
        "相关权限",
    ],

    "reporting_relationship": [
        "汇报关系",
    ],

    "collaboration_raw": [
        "工作协作关系",
    ],

    "work_environment": [
        "工作环境",
    ],

    "tools_and_equipment": [
        "使用工具设备",
    ],

    "required_documents": [
        "所需记录文档",
    ],

    "minimum_education": [
        "最低学历要求",
    ],

    "education_background": [
        "所需教育专业背景",
    ],

    "work_experience_raw": [
        "所需工作经验",
    ],

    "training_subjects": [
        "所需培训的科目",
    ],

    "competencies": [
        "所需胜任能力",
    ],

    "remarks": [
        "备注",
    ],
}

# 3. 基础文本清洗

def normalize_text(text: str | None) -> str:

    if not text:
        return ""

    text = (
        text
        .replace("\r\n", "\n")
        .replace("\r", "\n")
        .replace("\u3000", " ")
        .replace("\xa0", " ")
    )

    lines = []

    for line in text.split("\n"):

        line = re.sub(
            r"[ \t]+",
            " ",
            line
        )

        line = line.strip()

        if line:
            lines.append(line)

    return "\n".join(lines).strip()


def strip_list_prefix(text: str) -> str:

    text = text.strip()

    text = re.sub(
        r"^[•·●○▪■□◆◇▶►\-–—]+\s*",
        "",
        text,
    )

    # 数字编号
    text = re.sub(
        r"^\s*(?:"
        r"\d+\s*[\.、．]"
        r"|"
        r"[（(]\s*\d+\s*[）)]"
        r")\s*",
        "",
        text,
    )

    return text.strip()

# 4. Word 合并单元格处理

def get_unique_cell_texts(row) -> list[str]:
    """
    python-docx 在读取合并单元格时：row.cells
    可能会把同一个底层 XML cell 返回多次。
    cell._tc 判断是否是同一个真实单元格。
    """

    result: list[str] = []

    visited_cells: set[int] = set()

    for cell in row.cells:

        cell_id = id(cell._tc)

        if cell_id in visited_cells:
            continue

        visited_cells.add(cell_id)

        text = normalize_text(cell.text)

        result.append(text)

    return result

# 5. 字段提取工具

def extract_label_value(
    text: str,
    label: str,
) -> str | None:

    text = normalize_text(text)

    pattern = (
        rf"^\s*"
        rf"{re.escape(label)}"
        rf"\s*[：:]"
        rf"\s*(.*)$"
    )

    match = re.match(
        pattern,
        text,
        flags=re.S,
    )

    if not match:
        return None

    value = match.group(1).strip()

    return value or None


def extract_first(
    segment,
    aliases: list[str],
) -> str | None:

    for _, _, cells in segment:

        for text in cells:

            for label in aliases:

                value = extract_label_value(
                    text,
                    label,
                )

                if value is not None:
                    return value

    return None


# 6. 常用数据类型转换

def parse_int(
    value: str | None,
) -> int | None:

    if not value:
        return None

    match = re.search(
        r"\d+",
        value,
    )

    if not match:
        return None

    return int(match.group())


def parse_percentage(
    value: str | None,
) -> float | None:

    if not value:
        return None

    match = re.search(
        r"(\d+(?:\.\d+)?)\s*%",
        value,
    )

    if not match:
        return None

    return float(match.group(1))


# 7. 列表字段处理

def split_lines(
    value: str | None,
) -> list[str]:

    if not value:
        return []

    result = []

    for line in normalize_text(value).splitlines():

        line = strip_list_prefix(line)

        if line:
            result.append(line)

    return result


def split_delimited(
    value: str | None,
) -> list[str]:

    if not value:
        return []

    items = re.split(
        r"[\n，,；;、]+",
        normalize_text(value),
    )

    result = []

    for item in items:

        item = strip_list_prefix(item)

        if item:
            result.append(item)

    return result

# 8. 把整个 DOCX 转换为“行流”

def collect_row_stream(
    doc: Document,
):
    """
    把所有顶层 table 按顺序转换成：
    [(table_index, row_index, cells),...]
    """

    rows = []

    for table_index, table in enumerate(doc.tables):

        for row_index, row in enumerate(table.rows):

            cells = get_unique_cell_texts(row)

            rows.append(
                (
                    table_index,
                    row_index,
                    cells,
                )
            )

    return rows

# 9. 根据“岗位名称”切岗位

def find_job_title(
    cells: list[str],
) -> str | None:

    for text in cells:

        for alias in FIELD_ALIASES["job_title"]:

            value = extract_label_value(
                text,
                alias,
            )

            if value is not None:
                return value

    return None


def split_job_segments(rows):
    """
    遇到：岗位名称：XXX就认为一个新岗位开始。
    下一个岗位名称出现之前的所有行，都属于当前岗位。
    """

    segments = []

    current_segment = []

    for row in rows:

        _, _, cells = row

        title = find_job_title(cells)

        # 新岗位开始

        if title is not None:

            # 保存上一岗位
            if current_segment:
                segments.append(
                    current_segment
                )

            current_segment = [
                row
            ]

        # 当前岗位继续

        elif current_segment:

            current_segment.append(
                row
            )

    # 最后一个岗位
    if current_segment:

        segments.append(
            current_segment
        )

    return segments

# 10. 职责解析

def parse_responsibilities(
    segment,
) -> list[RawResponsibility]:
    """
    解析：
    职责描述：XXXX
    工作时间百分比：20%
    工作任务 | task1
    工作任务 | task2
    工作任务 | task3
    """

    responsibilities = []

    current: RawResponsibility | None = None

    for _, _, cells in segment:

        description = None
        percentage = None

        # 找职责描述 / 时间比例

        for text in cells:

            value = extract_label_value(
                text,
                "职责描述",
            )

            if value is not None:
                description = value

            value = extract_label_value(
                text,
                "工作时间百分比",
            )

            if value is not None:
                percentage = parse_percentage(
                    value
                )

        # 发现新的职责

        if description is not None:

            if current is not None:
                responsibilities.append(
                    current
                )

            current = RawResponsibility(
                sequence=len(
                    responsibilities
                ) + 1,
                description=description,
                time_percentage=percentage,
                tasks=[],
            )

            continue

        # 工作任务

        if current is not None:

            has_task_label = any(
                normalize_text(text)
                in {
                    "工作任务",
                    "工作 任务",
                }
                for text in cells
            )

            if not has_task_label:
                continue

            for text in cells:

                text = normalize_text(text)

                if not text:
                    continue

                if text in {
                    "工作任务",
                    "工作 任务",
                }:
                    continue

                if text.startswith(
                    "职责描述"
                ):
                    continue

                if text.startswith(
                    "工作时间百分比"
                ):
                    continue

                task = strip_list_prefix(
                    text
                )

                if (
                    task
                    and task not in current.tasks
                ):
                    current.tasks.append(
                        task
                    )

    # 保存最后一个职责
    if current is not None:

        responsibilities.append(
            current
        )

    return responsibilities

# 11. 工作协作关系解析

def parse_collaboration(
    raw: str | None,
) -> tuple[str | None, str | None]:

    if not raw:
        return None, None

    internal = None
    external = None

    for line in normalize_text(
        raw
    ).splitlines():

        value = extract_label_value(
            line,
            "内部协作部门",
        )

        if value is not None:
            internal = value

        value = extract_label_value(
            line,
            "外部协作单位",
        )

        if value is not None:
            external = value

    return internal, external

# 12. 原始字段保留

def extract_raw_fields(
    segment,
) -> dict[str, str]:

    result: dict[str, str] = {}

    for _, _, cells in segment:

        for text in cells:

            for aliases in FIELD_ALIASES.values():

                for label in aliases:

                    value = extract_label_value(
                        text,
                        label,
                    )

                    if value is not None:

                        result[label] = value

                        break

    return result

# 13. 数据质量检查

def build_warnings(
    job: RawJobDocument,
) -> list[str]:

    warnings = []

    if not job.job_summary:
        warnings.append(
            "缺少本职概述"
        )

    if not job.responsibilities:
        warnings.append(
            "未解析到职责与工作任务"
        )

    if (
        not job.qualification.minimum_education
    ):
        warnings.append(
            "缺少最低学历要求"
        )

    if (
        not job.qualification.work_experience_raw
    ):
        warnings.append(
            "缺少所需工作经验"
        )

    if (
        not job.qualification.competencies
    ):
        warnings.append(
            "缺少所需胜任能力"
        )

    return warnings

# 14. 解析一个岗位

def parse_job_segment(
    segment,
    source_file: str,
) -> RawJobDocument:

    start_table, start_row, _ = (
        segment[0]
    )

    end_table, end_row, _ = (
        segment[-1]
    )

    # 岗位名称

    title = extract_first(
        segment,
        FIELD_ALIASES[
            "job_title"
        ],
    )

    if not title:

        raise ValueError(
            "发现岗位起始位置，"
            "但没有解析到岗位名称："
            f"table={start_table}, "
            f"row={start_row}"
        )

    # 工作协作关系

    collaboration_raw = (
        extract_first(
            segment,
            FIELD_ALIASES[
                "collaboration_raw"
            ],
        )
    )

    (
        internal_collaboration,
        external_collaboration,
    ) = parse_collaboration(
        collaboration_raw
    )

    # 任职资格

    qualification = (
        RawQualification(

            minimum_education=
            extract_first(
                segment,
                FIELD_ALIASES[
                    "minimum_education"
                ],
            ),

            education_background=
            extract_first(
                segment,
                FIELD_ALIASES[
                    "education_background"
                ],
            ),

            work_experience_raw=
            extract_first(
                segment,
                FIELD_ALIASES[
                    "work_experience_raw"
                ],
            ),

            training_subjects=
            split_lines(
                extract_first(
                    segment,
                    FIELD_ALIASES[
                        "training_subjects"
                    ],
                )
            ),

            competencies=
            split_lines(
                extract_first(
                    segment,
                    FIELD_ALIASES[
                        "competencies"
                    ],
                )
            ),
        )
    )

    # 构造岗位

    job = RawJobDocument(

        # Source
        source_file=source_file,

        source_start_table=start_table,
        source_start_row=start_row,

        source_end_table=end_table,
        source_end_row=end_row,

        # Basic
        job_title=title,

        job_code=extract_first(
            segment,
            FIELD_ALIASES[
                "job_code"
            ],
        ),

        department=extract_first(
            segment,
            FIELD_ALIASES[
                "department"
            ],
        ),

        headcount=parse_int(
            extract_first(
                segment,
                FIELD_ALIASES[
                    "headcount"
                ],
            )
        ),

        direct_supervisor=
        extract_first(
            segment,
            FIELD_ALIASES[
                "direct_supervisor"
            ],
        ),

        direct_reports=
        split_delimited(
            extract_first(
                segment,
                FIELD_ALIASES[
                    "direct_reports"
                ],
            )
        ),

        job_family=extract_first(
            segment,
            FIELD_ALIASES[
                "job_family"
            ],
        ),

        team_size=parse_int(
            extract_first(
                segment,
                FIELD_ALIASES[
                    "team_size"
                ],
            )
        ),

        salary_grade_raw=
        extract_first(
            segment,
            FIELD_ALIASES[
                "salary_grade_raw"
            ],
        ),

        analysis_date_raw=
        extract_first(
            segment,
            FIELD_ALIASES[
                "analysis_date_raw"
            ],
        ),

        # Job content
        job_summary=extract_first(
            segment,
            FIELD_ALIASES[
                "job_summary"
            ],
        ),

        responsibilities=
        parse_responsibilities(
            segment
        ),

        permissions=
        split_lines(
            extract_first(
                segment,
                FIELD_ALIASES[
                    "permissions"
                ],
            )
        ),

        # Relationship
        reporting_relationship=
        extract_first(
            segment,
            FIELD_ALIASES[
                "reporting_relationship"
            ],
        ),

        internal_collaboration=
        internal_collaboration,

        external_collaboration=
        external_collaboration,

        # Environment
        work_environment=
        extract_first(
            segment,
            FIELD_ALIASES[
                "work_environment"
            ],
        ),

        tools_and_equipment=
        split_delimited(
            extract_first(
                segment,
                FIELD_ALIASES[
                    "tools_and_equipment"
                ],
            )
        ),

        required_documents=
        split_delimited(
            extract_first(
                segment,
                FIELD_ALIASES[
                    "required_documents"
                ],
            )
        ),

        # Qualification
        qualification=qualification,

        # Other
        remarks=extract_first(
            segment,
            FIELD_ALIASES[
                "remarks"
            ],
        ),

        raw_fields=
        extract_raw_fields(
            segment
        ),
    )

    # 数据质量检查

    job.warnings = (
        build_warnings(job)
    )

    return job

# 15. 主解析函数

def parse_jd_docx(
    input_path: str | Path,
) -> list[RawJobDocument]:

    input_path = Path(
        input_path
    )

    if not input_path.exists():

        raise FileNotFoundError(
            f"文件不存在：{input_path}"
        )

    if (
        input_path.suffix.lower()
        != ".docx"
    ):

        raise ValueError(
            "当前 Parser 只支持 .docx"
        )

    document = Document(
        str(input_path)
    )

    rows = collect_row_stream(
        document
    )

    segments = split_job_segments(
        rows
    )

    if not segments:

        raise ValueError(
            "没有发现任何岗位。"
            "请检查文档是否包含"
            "“岗位名称：XXX”字段。"
        )

    jobs = []

    for segment in segments:

        try:

            job = parse_job_segment(
                segment,
                input_path.name,
            )

            jobs.append(job)

        except Exception as exc:

            start_table, start_row, _ = (
                segment[0]
            )

            raise RuntimeError(
                "岗位解析失败："
                f"table={start_table}, "
                f"row={start_row}"
            ) from exc

    return jobs

# 16. 保存 JSON

def save_json(
    jobs: list[RawJobDocument],
    output_path: str | Path,
) -> None:

    output_path = Path(
        output_path
    )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    data = [
        job.model_dump()
        for job in jobs
    ]

    with output_path.open(
        "w",
        encoding="utf-8",
    ) as f:

        json.dump(
            data,
            f,
            ensure_ascii=False,
            indent=2,
        )

# 17. 主函数入口

def main():
    input_path = r"/jd预处理\岗位说明书.docx"
    output_path = "../jobs_raw.json"

    jobs = parse_jd_docx(input_path)

    save_json(
        jobs,
        output_path
    )

    print(f"共解析到 {len(jobs)} 个岗位")

    for index, job in enumerate(jobs, start=1):
        print(
            f"[{index}] "
            f"{job.job_title}"
            f" | 部门：{job.department or '-'}"
            f" | 职责数：{len(job.responsibilities)}"
            f" | warnings：{len(job.warnings)}"
        )

    print(f"结果已保存至：{output_path}")


if __name__ == "__main__":
    # main()
    from pathlib import Path
    import json

    project_root = Path(__file__).resolve().parents[1]
    input_path = project_root / "jobs_raw.json"

    INPUT_FILE = Path(input_path)
    print(f"INPUT_FILE: {INPUT_FILE}")

    with open(INPUT_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
        for i,item in enumerate(data):
            print(i,item["department"])